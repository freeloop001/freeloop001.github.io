# Word 文件处理示例
# 需要安装: pip install python-docx

try:
    from docx import Document
    import os

    # 创建 Word 文档
    doc = Document()
    doc.add_heading('Python 学习笔记', 0)

    doc.add_paragraph('这是一篇关于 Python 基础的文章')
    doc.add_paragraph('内容包括：变量、函数、类等')

    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = '姓名'
    table.cell(0, 1).text = '年龄'

    # 保存文件
    doc.save("notes.docx")
    print("Word 文档已创建: notes.docx")

    # 读取 Word 文档
    doc = Document("notes.docx")
    print("\n读取 Word 文档内容:")
    for para in doc.paragraphs:
        print(para.text)

    # 清理
    if os.path.exists("notes.docx"):
        os.remove("notes.docx")

except ImportError:
    print("python-docx 未安装: pip install python-docx")
